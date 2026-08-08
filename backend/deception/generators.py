"""
Fake-but-plausible artifact generators for the MIRAGE deception platform.
Nothing generated here is functional or valid.
"""

from __future__ import annotations

import json
import os
import random
import secrets
import string
import uuid
from datetime import datetime, timedelta, timezone

from dotenv import load_dotenv
from google import genai

from openai import AsyncOpenAI
from dotenv import load_dotenv
import os

load_dotenv()

client = AsyncOpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise RuntimeError("GEMINI_API_KEY not found in .env file")

client = genai.Client(api_key=api_key)


async def generate_fake_aws_key():
    prompt = """
Generate realistic-looking FAKE AWS credentials.

Rules:
- Never generate real credentials.
- Return ONLY valid JSON.
- No markdown.

Example:

{
  "access_key_id": "AKIAEXAMPLE12345678",
  "secret_access_key": "example_fake_secret_key"
}
"""

    response = await client.responses.create(
    model="gpt-5-mini",
    input=prompt,
    text={
        "format": {
            "type": "json_schema",
            "name": "aws_credentials",
            "schema": {
                "type": "object",
                "properties": {
                    "access_key_id": {"type": "string"},
                    "secret_access_key": {"type": "string"}
                },
                "required": [
                    "access_key_id",
                    "secret_access_key"
                ],
                "additionalProperties": False
            }
        }
    }
)

    text = response.output_text

    try:
        return json.loads(text)
    except Exception:
        return {
            "access_key_id": "AKIAFAKE1234567890",
            "secret_access_key": "mirage_fake_secret_key",
        }


_UPPER_NUM = string.ascii_uppercase + string.digits
_B64ISH = string.ascii_letters + string.digits + "+/"


def _rand(alphabet: str, length: int) -> str:
    return "".join(secrets.choice(alphabet) for _ in range(length))


def new_trigger_id() -> str:
    """Short URL-safe trigger ID."""
    return secrets.token_urlsafe(9).replace("-", "").replace("_", "")[:12]


def aws_credentials() -> dict:
    return {
        "aws_access_key_id": "AKIA" + _rand(_UPPER_NUM, 16),
        "aws_secret_access_key": _rand(_B64ISH, 40),
        "region": random.choice(
            [
                "us-east-1",
                "eu-west-1",
                "ap-south-1",
            ]
        ),
    }


def gcp_service_account(project="prod-data-pipeline") -> dict:
    return {
        "type": "service_account",
        "project_id": project,
        "private_key_id": uuid.uuid4().hex,
        "client_email": f"backup-agent@{project}.iam.gserviceaccount.com",
        "token_uri": "https://oauth2.googleapis.com/token",
    }


def azure_credentials() -> dict:
    return {
        "tenant_id": str(uuid.uuid4()),
        "client_id": str(uuid.uuid4()),
        "client_secret": _rand(
            string.ascii_letters + string.digits + "~.*-",
            40,
        ),
        "subscription_id": str(uuid.uuid4()),
    }


def db_connection_string(engine="postgresql") -> dict:
    user = random.choice(
        [
            "svc_reporting",
            "app_prod",
            "etl_worker",
            "dbadmin",
        ]
    )

    password = _rand(
        string.ascii_letters + string.digits,
        20,
    )

    host = random.choice(
        [
            "prod-db-01.internal",
            "10.20.4.17",
            "pg-cluster.corp.local",
        ]
    )

    port = {
        "postgresql": 5432,
        "mysql": 3306,
        "mongodb": 27017,
    }.get(engine, 5432)

    database = random.choice(
        [
            "customers",
            "billing",
            "payments",
            "hr_records",
        ]
    )

    return {
        "engine": engine,
        "username": user,
        "password": password,
        "host": host,
        "port": port,
        "database": database,
        "dsn": f"{engine}://{user}:{password}@{host}:{port}/{database}",
    }


def api_key(vendor=None):
    vendor = vendor or random.choice(
        [
            "stripe",
            "github",
            "slack",
            "sendgrid",
        ]
    )

    values = {
        "stripe": "sk_live_" + _rand(
            string.ascii_letters + string.digits,
            24,
        ),
        "github": "ghp_" + _rand(
            string.ascii_letters + string.digits,
            36,
        ),
        "slack": "xoxb-"
        + _rand(string.digits, 12)
        + "-"
        + _rand(string.ascii_letters + string.digits, 24),
        "sendgrid": "SG."
        + _rand(string.ascii_letters + string.digits, 22)
        + "."
        + _rand(string.ascii_letters + string.digits, 43),
    }

    return {
        "vendor": vendor,
        "api_key": values[vendor],
    }


def ssh_private_key():
    body = "\n".join(
        _rand(_B64ISH + "=", 64)
        for _ in range(14)
    )

    return {
        "key_type": "rsa",
        "comment": random.choice(
            [
                "deploy@jenkins",
                "root@bastion-01",
                "ansible@ctl",
            ]
        ),
        "private_key": (
            "-----BEGIN RSA PRIVATE KEY-----\n"
            + body
            + "\n-----END RSA PRIVATE KEY-----"
        ),
    }


def s3_bucket():
    name = random.choice(
        [
            "corp-db-backups-prod",
            "internal-payroll-exports",
            "customer-pii-archive",
            "terraform-state-prod",
        ]
    )

    return {
        "bucket": name,
        "arn": f"arn:aws:s3:::{name}",
        "url": f"https://{name}.s3.amazonaws.com/",
        "note": "Decoy bucket. Not a real resource.",
    }


def k8s_secret(namespace="production"):
    import base64

    creds = db_connection_string()
    enc = lambda s: base64.b64encode(s.encode()).decode()

    return {
        "apiVersion": "v1",
        "kind": "Secret",
        "metadata": {
            "name": "prod-db-credentials",
            "namespace": namespace,
        },
        "type": "Opaque",
        "data": {
            "username": enc(creds["username"]),
            "password": enc(creds["password"]),
            "host": enc(creds["host"]),
        },
    }


def fake_user():
    first = random.choice(
        [
            "marcus",
            "priya",
            "dana",
            "kenji",
            "elena",
        ]
    )

    last = random.choice(
        [
            "okoro",
            "raman",
            "whitfield",
            "sato",
            "vasquez",
        ]
    )

    return {
        "username": f"{first}.{last}",
        "email": f"{first}.{last}@corp.local",
        "display_name": f"{first.title()} {last.title()}",
        "title": random.choice(
            [
                "Backup Operator",
                "Finance Analyst",
                "Domain Admin (svc)",
            ]
        ),
        "password_hint": _rand(
            string.ascii_letters + string.digits,
            14,
        ),
        "created": (
            datetime.now(timezone.utc)
            - timedelta(days=random.randint(200, 900))
        ).isoformat(),
    }

import io
from docx import Document
from docx.shared import Inches
from PIL import Image

def fake_document(trigger_id: str, callback_url: str) -> dict:
    return {
        "filename": "Q3_Financials_Confidential.docx",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "tracking_mechanism": "Remote template / Linked object callback",
        "callback_url": callback_url,
    }


def generate_document_bytes(trigger_id: str, callback_url: str) -> bytes:
    doc = Document()
    doc.add_heading('Q3 Financial Projections & Mergers - CONFIDENTIAL', 0)
    doc.add_paragraph('This document contains proprietary financial data regarding the upcoming acquisition. Distribution is strictly limited to the executive team.')
    doc.add_paragraph('Accessing this document is monitored.')
    img = Image.new('RGB', (1, 1), color=(255, 255, 255))
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    doc.add_picture(img_bytes, width=Inches(0.01), height=Inches(0.01))
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.read()