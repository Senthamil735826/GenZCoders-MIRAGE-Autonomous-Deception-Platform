from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import os

class DocumentDeception:
    
    @staticmethod
    def create_fake_pdf(filepath, title="Q4 Financial Report"):
        c = canvas.Canvas(filepath, pagesize=letter)
        c.drawString(100, 750, title)
        c.drawString(100, 720, "CONFIDENTIAL - INTERNAL USE ONLY")
        c.drawString(100, 700, f"Generated: {datetime.now()}")
        c.drawString(100, 670, "Revenue: $4,250,000")
        c.drawString(100, 650, "Project Alpha Budget: $850,000")
        c.save()
        return filepath
    
    @staticmethod
    def create_fake_docx(filepath, title="Strategic Plan 2025"):
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph("CONFIDENTIAL DOCUMENT")
        doc.add_paragraph(f"Date: {datetime.now()}")
        doc.add_paragraph("Executive Summary: Expansion into APAC markets planned for Q2")
        doc.add_paragraph("Budget Allocation: $2.3M for new data center")
        doc.save(filepath)
        return filepath
    
    @staticmethod
    def create_fake_spreadsheet(filepath):
        import csv
        with open(filepath, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Employee", "SSN", "Salary", "Bank Account"])
            writer.writerow(["John Doe", "XXX-XX-1234", "$150,000", "****5678"])
            writer.writerow(["Jane Smith", "XXX-XX-5678", "$175,000", "****9012"])
        return filepath